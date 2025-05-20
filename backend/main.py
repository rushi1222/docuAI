# backend/main.py
from flask import Flask, request, jsonify
from flask_cors import CORS
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
from docx import Document
from docx.text.run import Run
import fitz, tiktoken, io, base64, os, re

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ model â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
load_dotenv()
llm = ChatOpenAI(
    model       = "gpt-4o",
    temperature = 0,
    api_key     = os.getenv("OPENAI_API_KEY"),
    max_retries = 2,
)

# â€”â€”â€” system prompt that guides every reply â€”â€”â€”â€”â€”â€”â€”â€”â€”â€”â€”â€”â€”â€”
SYSTEM_PROMPT = (
    "You are an intelligent assistant that helps users with a variety of "
    "queries â€” from general questions to legal/tax document replies. "
    "â–ª If the input is a **general question**, answer clearly and completely. "
    "â–ª If the input is an **official or legal document** (e.g. tax notice), "
    "produce a professional reply letter the user can submit. "
    "â–ª If the user asks for a **summary**, give a concise summary. "
    "Respond in plain language, with clear structure, no unnecessary AI disclaimers."
)

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ token / chunk helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# enc = tiktoken.encoding_for_model("gpt-4o")
# def tokens(t: str) -> int: return len(enc.encode(t))
def tokens(t: str) -> int: return len(t.split())


def chunk_text(t: str, max_tok=6000):
    words, cur, cur_tok = t.split(), [], 0
    for w in words:
        w_tok = tokens(w + " ")
        if cur_tok + w_tok > max_tok:
            yield " ".join(cur);  cur, cur_tok = [], 0
        cur.append(w); cur_tok += w_tok
    if cur: yield " ".join(cur)

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ file â†’ text helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def text_from_docx(raw: bytes):
    doc = Document(io.BytesIO(raw))
    return "\n".join(p.text for p in doc.paragraphs)

def text_from_pdf(raw: bytes):
    pdf = fitz.open(stream=raw, filetype="pdf")
    return "\n".join(p.get_text() for p in pdf)

def extract_text(b64: str, name: str):
    data = base64.b64decode(b64)
    ext  = os.path.splitext(name.lower())[1]
    if ext == ".pdf":  return text_from_pdf(data)
    if ext == ".docx": return text_from_docx(data)
    return data.decode(errors="ignore")

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ long-text first-turn summary â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def summarise_big(txt: str) -> str:
    parts = [
        llm.invoke(f"Summarise clearly:\n\n{c}\n\nSummary:").content.strip()
        for c in chunk_text(txt)
    ]
    combined = "\n\n".join(parts)
    if tokens(combined) > 6000:
        combined = summarise_big(combined)          # recurse once
    return llm.invoke(
        "Merge and refine these summaries into one concise overview:\n\n"
        f"{combined}\n\nFinal summary (plain text, no markdown):"
    ).content.strip()

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ Markdown â†’ DOCX (bold + simple lists) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_BOLD = re.compile(r'\*\*(.*?)\*\*')

def add_md_para(doc: Document, md: str):
    for line in md.splitlines():
        p = doc.add_paragraph()

        # bullet / numbered list detection
        if re.match(r'^\s*[\-\*]\s+', line):
            p.add_run("â€¢ "); line = re.sub(r'^\s*[\-\*]\s+', '', line)
        elif re.match(r'^\s*\d+\.\s+', line):
            p.add_run("â–ª "); line = re.sub(r'^\s*\d+\.\s+', '', line)

        pos = 0
        for m in _BOLD.finditer(line):
            if m.start() > pos:
                p.add_run(line[pos:m.start()])
            bold: Run = p.add_run(m.group(1)); bold.bold = True
            pos = m.end()
        if pos < len(line):
            p.add_run(line[pos:])

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ Flask app & routes â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
app = Flask(__name__)
CORS(app)

@app.post("/chat")
def chat():
    p    = request.json or {}
    hist = p.get("history", [])
    text = (p.get("text") or "").strip()

    # if first turn includes a file â†’ extract + summarise
    if p.get("fileB64") and p.get("fileName"):
        try:
            text = extract_text(p["fileB64"], p["fileName"])
        except Exception as e:
            return jsonify({"error": f"file error: {e}"}), 400
        text = summarise_big(text)

    if not text:
        return jsonify({"error": "No input text"}), 400

    # build conversation with system prompt up front
    messages = [{"role": "system", "content": SYSTEM_PROMPT}, *hist,
                {"role": "user",   "content": text}]

    assistant = llm.invoke(messages).content.strip()
    return jsonify({"assistant": assistant})

@app.post("/download")
def download():
    j     = request.json or {}
    md    = j.get("content","").strip()
    fname = j.get("fileName","reply_output.docx").rsplit(".",1)[0] + ".docx"
    if not md:
        return jsonify({"error":"No content"}), 400

    doc = Document()
    for para in md.split("\n\n"):        # preserve blank lines
        add_md_para(doc, para)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return jsonify({
        "b64": base64.b64encode(buf.read()).decode(),
        "fileName": fname
    })

# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€ run server â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
if __name__ == "__main__":
    print("ðŸš€  backend running  http://0.0.0.0:8000")
    app.run(host="0.0.0.0", port=8000)
